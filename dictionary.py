# from docx import document

# docx_file = 'dictionary.docx'
# txt_file = 'dictionary.txt'

# docx = document.process(docx_file)

# for para in docx.paragraphs : 
#     text + = para.text +"\n"

  
# file=open(txt_file, 'w', encoding='utf-8')
# file.write(docx)
# file.close()

from docx import Document

def convert_docx_to_txt(docx_file, txt_file):
    # Load the .docx file
    doc = Document(docx_file)
    
    # Create a list to hold the paragraphs
    text = []

    # Iterate through each paragraph in the document
    for para in doc.paragraphs:
        text.append(para.text)

    # Join all paragraphs into a single string with new lines
    full_text = '\n'.join(text)

    # Write the text to the output file
    with open(txt_file, 'w', encoding='utf-8') as file:
        file.write(full_text)

# File paths
docx_file = 'dictionary.docx'  # Replace with your input file path
txt_file = 'dictionary.txt'      # Replace with your desired output file path

# Convert the .docx file to .txt
convert_docx_to_txt(docx_file, txt_file)

print("Conversion complete. Text file saved as:", txt_file)